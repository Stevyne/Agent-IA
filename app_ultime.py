"""
🌟🤖👁️🎙️📝 AGENT ULTIME — Tout-en-un 100% local

Interface Streamlit combinant :
  • Chat texte + vision (Qwen 2.5 / LLaVA)
  • Outils : calcul, heure, recherche documents, création de fichiers
  • Mémoire persistante (JSON)
  • Voix synthétique (TTS) : pyttsx3 (offline) ou Piper (naturelle, offline)
  • Micro (STT) : Faster-Whisper (offline) ou SpeechRecognition (Google)
  • Export documents : TXT, MD, CSV, JSON, DOCX, PDF

INSTALLATION :
  pip install streamlit pillow ollama python-docx fpdf2
  # Optionnel voix :
  pip install pyttsx3
  # Optionnel micro offline :
  pip install faster-whisper pyaudio SpeechRecognition
  # Optionnel RAG documents :
  pip install chromadb pypdf python-docx

OLLAMA :
  ollama pull qwen2.5:3b
  ollama pull llava-phi3   # ou moondream

Piper-TTS (optionnel, voix naturelle offline) :
  - Téléchargez Piper + une voix fr dans models/piper/
  - Voir GUIDE_VOIX_OFFLINE.md

USAGE :
  streamlit run app_ultime.py
"""

import streamlit as st
import ollama
import os
import sys
import json
import csv
import base64
import re
import tempfile
import subprocess
import platform
import threading
import wave
import time
from datetime import datetime
from io import BytesIO
from translations import TRANSLATIONS, LANG_NAMES

# ========================= CONFIGURATION =========================
MODEL_TEXT = "qwen2.5:3b"
MODEL_VISION = "llava-phi3"
MEMORY_FILE = "memory_ultime.json"
OUTPUTS_DIR = "outputs"
MAX_TEXT_HISTORY = 20
MAX_VISION_HISTORY = 10
MAX_IMAGE_SIZE = 1024
RAG_DB_PATH = "chroma_db_docs"
RAG_COLLECTION = "bibliotheque"

# TTS
USE_PIPER = False              # True pour Piper (naturelle), False pour pyttsx3 (simple)
PIPER_DIR = "models/piper"
PIPER_MODEL = None             # None = auto-detect
PIPER_LENGTH_SCALE = 1.0
VOIX_VITESSE = 170             # pour pyttsx3

# STT
USE_FASTER_WHISPER = True      # True = offline (local), False = Google SpeechRecognition
WHISPER_MODEL = "tiny"         # tiny (39 Mo), base (74 Mo), small (244 Mo)
WHISPER_DEVICE = "cpu"
WHISPER_COMPUTE_TYPE = "int8"
# ================================================================


# --- INITIALISATION SESSION_STATE ---
if "lang" not in st.session_state:
    st.session_state.lang = "fr" 

# --- FONCTION MAGIQUE DE TRADUCTION t() ---
def t(key: str) -> str:
    lang = st.session_state.lang
    return TRANSLATIONS.get(lang, TRANSLATIONS["fr"]).get(key, key)


# --- Détection imports optionnels ---
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import chromadb
    from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
    HAS_CHROMADB = True
except ImportError:
    HAS_CHROMADB = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

try:
    import pyttsx3
    HAS_PYTTSX3 = True
except ImportError:
    HAS_PYTTSX3 = False

try:
    import speech_recognition as sr
    HAS_SR = True
except ImportError:
    HAS_SR = False

try:
    from faster_whisper import WhisperModel
    HAS_WHISPER = True
except ImportError:
    HAS_WHISPER = False

try:
    import pyaudio
    HAS_PYAUDIO = True
except ImportError:
    HAS_PYAUDIO = False

try:
    from duckduckgo_search import DDGS
    HAS_DUCKDUCKGO = True
except ImportError:
    HAS_DUCKDUCKGO = False


# ========================= PERSONNALITÉS =========================

def load_personnalites():
    """Charge les personnalités depuis personnalites.json. Retourne une liste de dicts."""
    default_personnalites = [
        {"nom": "🤖 Assistant", "description": "Neutre et serviable", "system_prompt": "Tu es un assistant IA serviable, clair et concis."},
        {"nom": "👨‍🏫 Professeur", "description": "Patient et pédagogique", "system_prompt": "Tu es un professeur patient qui explique étape par étape avec des analogies."},
        {"nom": "👨‍💻 Développeur", "description": "Direct et technique", "system_prompt": "Tu es un développeur senior. Réponses directes, techniques, avec du code."},
        {"nom": "🏋️ Coach", "description": "Motivant et concis", "system_prompt": "Tu es un coach motivant. Réponses courtes, positives et actionnables."},
        {"nom": "🎨 Poète", "description": "Créatif et imagé", "system_prompt": "Tu es un poète créatif. Réponses avec métaphores et style littéraire."},
        {"nom": "😏 Sarcastic", "description": "Ironique mais utile", "system_prompt": "Tu es sarcastique et ironique, mais toujours utile et précis."},
        {"nom": "🧪 Scientifique", "description": "Rigoureux et sceptique", "system_prompt": "Tu es un scientifique rigoureux. Faits nuancés, preuves, scepticisme."},
        {"nom": "👶 Explicateur (5 ans)", "description": "Ultra simple", "system_prompt": "Tu expliques comme à un enfant de 5 ans. Mots simples, analogies ludiques."},
    ]
    if not os.path.exists("personnalites.json"):
        return default_personnalites
    try:
        with open("personnalites.json", "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict) and "personnalites" in data:
                return data["personnalites"]
            if isinstance(data, list):
                return data
    except Exception:
        pass
    return default_personnalites


def detecter_personnalite_auto(query: str, personnalites: list) -> int:
    """
    Détecte automatiquement la personnalité la plus appropriée selon la requête.
    Utilise une détection par mots-clés rapide (regex). Retourne l'index.
    """
    q = query.lower()
    # Index : mots-clés regex associés
    mots_cles = {
        1: [r"\bapprends?\b", r"\bexplique\b", r"\bcours\b", r"\bleçon\b", r"\btutoriel\b",
            r"\bpédagogique\b", r"\benseigne\b", r"\bétape\b", r"\bcomprendre\b",
            r"\bcomment ça marche\b", r"\bpourquoi\b", r"\bnotion\b", r"\bconcept\b"],
        2: [r"\bcode\b", r"\bcoder\b", r"\bprogramme\b", r"\bprogrammation\b", r"\bpython\b",
            r"\bjavascript\b", r"\bhtml\b", r"\bcss\b", r"\bbug\b", r"\bdébug\b", r"\bfonction\b",
            r"\bscript\b", r"\bdéveloppe\b", r"\bapi\b", r"\bbibliothèque\b", r"\blibrary\b",
            r"\bdéveloppeur\b", r"\bgit\b", r"\bgithub\b", r"\bdocker\b"],
        3: [r"\bmotive\b", r"\bmotivation\b", r"\bdéprime\b", r"\btriste\b", r"\bobjectif\b",
            r"\bbut\b", r"\bréussir\b", r"\bcourage\b", r"\baide[- ]moi\b", r"\bconseil\b",
            r"\bexercice\b", r"\bsport\b", r"\bpersévère\b", r"\babandonne\b"],
        4: [r"\bpoème\b", r"\bpoésie\b", r"\bpoétique\b", r"\brime\b", r"\brimes\b",
            r"\bécrit\b", r"\bstyle littéraire\b", r"\bmétaphore\b", r"\bimager\b",
            r"\bromantique\b", r"\bâme\b", r"\bvers\b", r"\bprose\b"],
        5: [r"\bdrôle\b", r"\bmarrant\b", r"\bhumour\b", r"\brigole\b", r"\bblague\b",
            r"\bamuse\b", r"\bsarcastique\b", r"\bironique\b", r"\bmoque\b", r"\btaquin\b",
            r"\bhilarant\b", r"\bfun\b"],
        6: [r"\bscience\b", r"\bétude\b", r"\bexpérience\b", r"\bpreuve\b", r"\bdonnées\b",
            r"\bstatistique\b", r"\bhypothèse\b", r"\bthéorie\b", r"\brecherche\b",
            r"\bacadémique\b", r"\banalyse rigoureuse\b", r"\bscientifique\b", r"\bpeer[- ]review\b"],
        7: [r"\bhistoire\b", r"\bhistorique\b", r"\bpassé\b", r"\bantique\b", r"\bsiècle\b",
            r"\bmoyen[- ]?âge\b", r"\brévolution\b", r"\bépoque\b", r"\bguerre\b",
            r"\bcivilisation\b", r"\bantiquité\b", r"\bempire\b", r"\bévénement historique\b"],
        8: [r"\bsimple\b", r"\bbébé\b", r"\benfant\b", r"\bfacile\b", r"\bpour les nuls\b",
            r"\bdummy\b", r"\bcomprens?ion\b", r"\bc'est quoi\b", r"\bdéfinition simple\b",
            r"\bcomme si j'avais 5 ans\b", r"\bcomme à un enfant\b"],
    }

    scores = {i: 0 for i in range(len(personnalites))}
    for idx, patterns in mots_cles.items():
        if idx >= len(personnalites):
            continue
        for pat in patterns:
            if re.search(pat, q):
                scores[idx] += 1

    best = max(scores, key=scores.get)
    if scores[best] > 0:
        return best
    return 0  # Assistant par défaut


# ========================= 0. VÉRIFICATION INTERNET =========================

def check_internet(timeout: float = 3.0) -> bool:
    """
    Vérifie si une connexion Internet est disponible.
    Essaie de contacter duckduckgo.com (port 443) pour la recherche web.
    """
    import socket
    try:
        host = socket.gethostbyname("duckduckgo.com")
        s = socket.create_connection((host, 443), timeout=timeout)
        s.close()
        return True
    except Exception:
        return False


# ========================= 1. TTS (SYNTHÈSE VOCALE) =========================

class TTSManager:
    """Gère la synthèse vocale : pyttsx3 (simple) ou Piper (naturelle)."""

    def __init__(self, use_piper: bool = USE_PIPER):
        self.use_piper = use_piper
        self.ok = False
        self._pyttsx_engine = None
        self._piper_bin = None
        self._piper_model = None

        if self.use_piper:
            self._init_piper()
        else:
            self._init_pyttsx3()

    def _init_pyttsx3(self):
        if not HAS_PYTTSX3:
            return
        try:
            self._pyttsx_engine = pyttsx3.init()
            self._pyttsx_engine.setProperty("rate", VOIX_VITESSE)
            voices = self._pyttsx_engine.getProperty("voices")
            for v in voices:
                if "french" in v.name.lower() or "franc" in v.name.lower():
                    self._pyttsx_engine.setProperty("voice", v.id)
                    break
            self.ok = True
        except Exception as e:
            st.sidebar.warning(f"TTS pyttsx3 erreur : {e}")

    def _find_piper(self):
        system = platform.system()
        name = "piper.exe" if system == "Windows" else "piper"
        local = os.path.join(PIPER_DIR, name)
        if os.path.isfile(local):
            return os.path.abspath(local)
        for path in os.environ.get("PATH", "").split(os.pathsep):
            candidate = os.path.join(path, name)
            if os.path.isfile(candidate):
                return candidate
        return None

    def _find_piper_model(self):
        if PIPER_MODEL is not None:
            candidate = os.path.join(PIPER_DIR, PIPER_MODEL)
            if os.path.isfile(candidate):
                return os.path.abspath(candidate)
        if not os.path.isdir(PIPER_DIR):
            return None
        for f in os.listdir(PIPER_DIR):
            if f.endswith(".onnx") and not f.endswith(".onnx.json"):
                return os.path.abspath(os.path.join(PIPER_DIR, f))
        return None

    def _init_piper(self):
        if not HAS_PYAUDIO:
            return
        self._piper_bin = self._find_piper()
        self._piper_model = self._find_piper_model()
        if self._piper_bin and self._piper_model:
            self.ok = True
            st.sidebar.success(f"Piper TTS prêt : {os.path.basename(self._piper_model)}")
        else:
            st.sidebar.warning("Piper TTS non trouvé. Basculez vers pyttsx3 ou configurez models/piper/.")

    def _generate_piper_wav(self, text: str) -> str | None:
        fd, wav_path = tempfile.mkstemp(suffix=".wav")
        os.close(fd)
        cmd = [self._piper_bin, "--model", self._piper_model, "--output_file", wav_path]
        if PIPER_LENGTH_SCALE != 1.0:
            cmd += ["--length_scale", str(PIPER_LENGTH_SCALE)]
        try:
            result = subprocess.run(cmd, input=text.encode("utf-8"), capture_output=True, shell=False)
            if result.returncode != 0:
                os.remove(wav_path)
                return None
            return wav_path
        except Exception:
            if os.path.exists(wav_path):
                os.remove(wav_path)
            return None

    def _play_wav(self, wav_path: str):
        if not HAS_PYAUDIO:
            return
        try:
            wf = wave.open(wav_path, "rb")
            p = pyaudio.PyAudio()
            stream = p.open(format=p.get_format_from_width(wf.getsampwidth()),
                            channels=wf.getnchannels(),
                            rate=wf.getframerate(),
                            output=True)
            data = wf.readframes(1024)
            while data:
                stream.write(data)
                data = wf.readframes(1024)
            stream.stop_stream()
            stream.close()
            p.terminate()
            wf.close()
        except Exception:
            pass
        finally:
            if os.path.exists(wav_path):
                os.remove(wav_path)

    def _clean_text(self, text: str) -> str:
        clean = text
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
        clean = re.sub(r"`(.+?)`", r"\1", clean)
        clean = re.sub(r"#+ ", "", clean)
        clean = clean.replace("\n", " ")
        return clean

    def speak(self, text: str):
        if not self.ok or not text:
            return
        clean = self._clean_text(text)
        if self.use_piper and self._piper_bin and self._piper_model:
            wav = self._generate_piper_wav(clean)
            if wav:
                self._play_wav(wav)
        elif self._pyttsx_engine:
            try:
                self._pyttsx_engine.say(clean)
                self._pyttsx_engine.runAndWait()
            except Exception:
                pass


# ========================= 2. STT (RECONNAISSANCE VOCALE) =========================

class STTManager:
    """Gère le micro : Faster-Whisper (offline) ou SpeechRecognition (Google)."""

    def __init__(self, use_whisper: bool = USE_FASTER_WHISPER):
        self.use_whisper = use_whisper
        self.ok = False
        self._whisper_model = None
        self._recognizer = None
        self._microphone = None

        if self.use_whisper and HAS_WHISPER and HAS_SR and HAS_PYAUDIO:
            self._init_whisper()
        elif not self.use_whisper and HAS_SR and HAS_PYAUDIO:
            self._init_google()

    def _init_whisper(self):
        try:
            self._whisper_model = WhisperModel(WHISPER_MODEL, device=WHISPER_DEVICE, compute_type=WHISPER_COMPUTE_TYPE)
            self._recognizer = sr.Recognizer()
            self._microphone = sr.Microphone()
            with self._microphone as source:
                self._recognizer.adjust_for_ambient_noise(source, duration=1)
            self.ok = True
        except Exception as e:
            st.sidebar.warning(f"Whisper STT erreur : {e}")

    def _init_google(self):
        try:
            self._recognizer = sr.Recognizer()
            self._microphone = sr.Microphone()
            with self._microphone as source:
                self._recognizer.adjust_for_ambient_noise(source, duration=1)
            self.ok = True
        except Exception as e:
            st.sidebar.warning(f"Google STT erreur : {e}")

    def listen(self) -> str | None:
        if not self.ok:
            return None
        try:
            with self._microphone as source:
                audio = self._recognizer.listen(source, timeout=5, phrase_time_limit=10)
        except Exception:
            return None

        try:
            if self.use_whisper and self._whisper_model:
                fd, wav_path = tempfile.mkstemp(suffix=".wav")
                os.close(fd)
                try:
                    with open(wav_path, "wb") as f:
                        f.write(audio.get_wav_data())
                    segments, _ = self._whisper_model.transcribe(wav_path, language="fr", condition_on_previous_text=False)
                    text = " ".join([seg.text.strip() for seg in segments]).strip()
                    return text if text else None
                finally:
                    if os.path.exists(wav_path):
                        os.remove(wav_path)
            else:
                return self._recognizer.recognize_google(audio, language="fr-FR")
        except Exception:
            return None


# ========================= 3. SÉCURITÉ / OUTPUTS =========================

def _sanitize_filename(name: str) -> str:
    if not name:
        return "document.txt"
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = name.replace("..", "_")
    name = name.lstrip("/\\")
    return name


def _ensure_outputs_dir():
    if not os.path.exists(OUTPUTS_DIR):
        os.makedirs(OUTPUTS_DIR)


def _get_safe_path(name: str) -> str:
    _ensure_outputs_dir()
    safe = _sanitize_filename(name)
    return os.path.join(OUTPUTS_DIR, safe)


# ========================= 4. OUTILS TEXTE =========================

def calculer(expression: str) -> str:
    if not re.match(r'^[0-9+\-*/.() ]+$', expression):
        return "Erreur : caractères non autorisés."
    try:
        return str(eval(expression))
    except Exception as e:
        return f"Erreur de calcul : {e}"


def heure_actuelle() -> str:
    return datetime.now().strftime("%A %d %B %Y, %H:%M:%S")


# --- AGENDA / TODO ---

AGENDA_FILE = "agenda.json"


def _charger_agenda():
    if not os.path.exists(AGENDA_FILE):
        return {"taches": []}
    try:
        with open(AGENDA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict) and "taches" in data:
                return data
            return {"taches": []}
    except Exception:
        return {"taches": []}


def _sauver_agenda(data):
    try:
        with open(AGENDA_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        pass


def _generer_id_tache():
    import random
    return f"t_{random.randint(1000, 9999)}_{random.randint(1000, 9999)}"


def ajouter_tache(nom: str, description: str = None, date: str = None, priorite: str = "normale") -> str:
    """Ajoute une tâche à l'agenda."""
    if not nom or not str(nom).strip():
        return "Erreur : le nom de la tâche ne peut pas être vide."
    data = _charger_agenda()
    tache = {
        "id": _generer_id_tache(),
        "nom": str(nom).strip(),
        "description": str(description).strip() if description else None,
        "date": str(date).strip() if date else None,
        "priorite": str(priorite).lower().strip() if priorite else "normale",
        "fait": False,
        "date_creation": datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    }
    data["taches"].append(tache)
    _sauver_agenda(data)
    return f"✅ Tâche ajoutée : '{tache['nom']}' (ID: {tache['id']}, priorité: {tache['priorite']})"


def lister_taches(filtre_date: str = None, filtre_statut: str = None, limite: int = 20) -> str:
    """Liste les tâches de l'agenda, avec filtres optionnels."""
    data = _charger_agenda()
    taches = data.get("taches", [])
    if not taches:
        return "📋 Votre agenda est vide."

    taches_filtrees = []
    for t in taches:
        if filtre_date and t.get("date") != filtre_date:
            continue
        if filtre_statut:
            fs = str(filtre_statut).lower().strip()
            if fs in ("fait", "done", "terminé", "terminée") and not t.get("fait"):
                continue
            if fs in ("afaire", "à faire", "todo", "non fait", "non fait") and t.get("fait"):
                continue
        taches_filtrees.append(t)

    taches_filtrees = taches_filtrees[:int(limite)] if limite else taches_filtrees

    if not taches_filtrees:
        return "📋 Aucune tâche ne correspond à vos critères."

    lignes = []
    for i, t in enumerate(taches_filtrees, 1):
        statut = "✅" if t.get("fait") else "⬜"
        date_str = f" ({t['date']})" if t.get("date") else ""
        prio = f" [{t.get('priorite','normale')}]" if t.get('priorite') and t.get('priorite') != 'normale' else ""
        desc = f" — {t['description']}" if t.get('description') else ""
        lignes.append(f"{i}. {statut} {t['nom']}{date_str}{prio}{desc} (ID: {t['id']})")

    return "📋 Vos tâches :\n" + "\n".join(lignes)


def marquer_fait(tache_id: str) -> str:
    """Marque une tâche comme faite."""
    data = _charger_agenda()
    for t in data.get("taches", []):
        if t.get("id") == tache_id:
            t["fait"] = True
            _sauver_agenda(data)
            return f"✅ Tâche '{t['nom']}' marquée comme faite."
    # Recherche partielle si l'ID exact ne correspond pas (l'IA peut se tromper)
    for t in data.get("taches", []):
        if tache_id.lower() in t.get("nom", "").lower() or tache_id.lower() in t.get("id", "").lower():
            t["fait"] = True
            _sauver_agenda(data)
            return f"✅ Tâche '{t['nom']}' (ID: {t['id']}) marquée comme faite."
    return f"❌ Tâche '{tache_id}' introuvable."


def supprimer_tache(tache_id: str) -> str:
    """Supprime une tâche de l'agenda."""
    data = _charger_agenda()
    original_len = len(data.get("taches", []))
    data["taches"] = [t for t in data.get("taches", []) if t.get("id") != tache_id]
    if len(data["taches"]) < original_len:
        _sauver_agenda(data)
        return f"🗑️ Tâche supprimée."
    # Recherche partielle
    nouvelles = [t for t in data.get("taches", []) if tache_id.lower() not in t.get("nom", "").lower() and tache_id.lower() not in t.get("id", "").lower()]
    if len(nouvelles) < len(data["taches"]):
        data["taches"] = nouvelles
        _sauver_agenda(data)
        return f"🗑️ Tâche correspondant à '{tache_id}' supprimée."
    return f"❌ Tâche '{tache_id}' introuvable."


# --- RECHERCHE WEB (DuckDuckGo) ---

def rechercher_web(requete: str, n_resultats: int = 3) -> str:
    """
    Recherche sur le web via DuckDuckGo (gratuit, sans clé API).
    Retourne un texte avec les titres et snippets des résultats.
    """
    if not HAS_DUCKDUCKGO:
        return "Erreur : bibliothèque 'duckduckgo-search' non installée. Tapez : pip install duckduckgo-search"

    if not check_internet():
        return "Vous semblez hors ligne. Je ne peux pas effectuer de recherche web pour l'instant. Vérifiez votre connexion Internet."

    try:
        n = int(n_resultats)
        if n < 1:
            n = 1
        if n > 10:
            n = 10
    except Exception:
        n = 3

    try:
        with DDGS() as ddgs:
            results = ddgs.text(requete, max_results=n)
            if not results:
                return "Aucun résultat trouvé pour cette recherche."

            lines = []
            for i, r in enumerate(results, 1):
                title = r.get("title", "Sans titre")
                body = r.get("body", "")
                href = r.get("href", "")
                lines.append(f"{i}. {title}\n{body}\nSource: {href}")
            return "\n\n".join(lines)
    except Exception as e:
        return f"Erreur lors de la recherche web : {e}"


# --- RAG ---
_rag_collection = None

def init_rag():
    global _rag_collection
    if not HAS_CHROMADB:
        return False
    try:
        client = chromadb.PersistentClient(path=RAG_DB_PATH)
        _rag_collection = client.get_collection(name=RAG_COLLECTION)
        return True
    except Exception:
        _rag_collection = None
        return False


def rechercher_documents(requete: str, n_resultats: int = 5) -> str:
    if _rag_collection is None:
        return "Base de documents non initialisée."
    try:
        n_resultats = int(n_resultats)
        if n_resultats < 1:
            n_resultats = 1
        if n_resultats > 10:
            n_resultats = 10
    except Exception:
        n_resultats = 5
    try:
        results = _rag_collection.query(query_texts=[requete], n_results=n_resultats)
    except Exception as e:
        return f"Erreur recherche : {e}"
    if not results or not results["documents"] or not results["documents"][0]:
        return "Aucun passage pertinent trouvé."
    passages = []
    for i, doc in enumerate(results["documents"][0]):
        meta = results["metadatas"][0][i] if results.get("metadatas") else {}
        source = meta.get("source", "inconnu")
        doc_type = meta.get("type", "?")
        page = meta.get("page", "?")
        if doc_type == "pdf":
            loc = f"page {page}"
        elif doc_type == "docx":
            loc = f"paragraphe {page}"
        elif doc_type == "txt":
            loc = f"bloc {page}"
        else:
            loc = f"position {page}"
        passages.append(f"---\nSource: {source} ({loc})\n{doc}")
    return "\n".join(passages)


# --- BAC À SABLE PYTHON ---

def executer_python(code: str, nom_fichier_sortie: str = None) -> str:
    """
    Exécute du code Python dans un bac à sable sécurisé.
    Interdit : os.system, subprocess, socket, urllib, eval, exec, fichiers hors sandbox.
    Retourne stdout, stderr, et les fichiers (graphiques) créés.
    """
    mots_interdits = [
        "import os", "from os", "import subprocess", "from subprocess",
        "import sys", "import socket", "import urllib", "import requests",
        "import ftplib", "import telnetlib", "import smtplib", "import email",
        "import http", "import webbrowser", "import pip", "import importlib",
        "eval(", "exec(", "__import__(", "compile(", "open(", "file(",
        "subprocess.run", "subprocess.call", "subprocess.Popen", "subprocess.check_output",
        "os.system", "os.popen", "os.remove", "os.rmdir", "os.unlink", "os.rename", "os.replace",
        "os.chmod", "os.mkdir", "os.makedirs", "os.walk", "os.listdir", "os.scandir",
        "shutil.rmtree", "shutil.copy", "shutil.move", "shutil.copytree",
        "breakpoint(", "input(", "raw_input(", "exit(", "quit(",
    ]
    code_lower = code.lower()
    for mot in mots_interdits:
        if mot.lower() in code_lower:
            return f"🛡️ **ERREUR SÉCURITÉ** : Le code contient un élément interdit : `{mot}`. Pour votre protection, l'accès au système, au réseau et aux fichiers hors du bac à sable est bloqué."

    sandbox_dir = os.path.join(OUTPUTS_DIR, "sandbox")
    os.makedirs(sandbox_dir, exist_ok=True)

    # Wrapper qui capture stdout et sécurise matplotlib
    wrapper = f'''import sys, json, os, io, traceback
import matplotlib
matplotlib.use('Agg')
try:
    import matplotlib.pyplot as plt
except Exception:
    plt = None
try:
    import numpy as np
except Exception:
    np = None
try:
    import pandas as pd
except Exception:
    pd = None

sandbox_dir = {repr(sandbox_dir)}

# Désactivation de fonctions dangereuses
import subprocess, os as os_module
subprocess.run = lambda *a, **k: (_ for _ in ()).throw(RuntimeError("subprocess interdit"))
subprocess.call = lambda *a, **k: (_ for _ in ()).throw(RuntimeError("subprocess interdit"))
subprocess.Popen = lambda *a, **k: (_ for _ in ()).throw(RuntimeError("subprocess interdit"))
os_module.system = lambda *a: (_ for _ in ()).throw(RuntimeError("os.system interdit"))
os_module.popen = lambda *a, **k: (_ for _ in ()).throw(RuntimeError("os.popen interdit"))
os_module.remove = lambda *a: (_ for _ in ()).throw(RuntimeError("os.remove interdit"))
os_module.rmdir = lambda *a: (_ for _ in ()).throw(RuntimeError("os.rmdir interdit"))
os_module.unlink = lambda *a: (_ for _ in ()).throw(RuntimeError("os.unlink interdit"))
os_module.rename = lambda *a, **k: (_ for _ in ()).throw(RuntimeError("os.rename interdit"))
os_module.replace = lambda *a, **k: (_ for _ in ()).throw(RuntimeError("os.replace interdit"))
os_module.chmod = lambda *a, **k: (_ for _ in ()).throw(RuntimeError("os.chmod interdit"))

stdout_buf = io.StringIO()
stderr_buf = io.StringIO()
old_out, old_err = sys.stdout, sys.stderr
sys.stdout, sys.stderr = stdout_buf, stderr_buf

result = {{"success": False, "output": "", "error": "", "files": []}}

try:
{chr(10).join("    " + line for line in code.split(chr(10)))}

    if plt is not None:
        for i in plt.get_fignums():
            fig = plt.figure(i)
            if {repr(nom_fichier_sortie)}:
                fname = os.path.join(sandbox_dir, {repr(nom_fichier_sortie)})
            else:
                fname = os.path.join(sandbox_dir, f"figure_{{i}}.png")
            fig.savefig(fname, dpi=150, bbox_inches='tight')
            result["files"].append(fname)
        plt.close('all')
    result["success"] = True
except Exception as e:
    result["error"] = traceback.format_exc()

sys.stdout, sys.stderr = old_out, old_err
result["output"] = stdout_buf.getvalue()
result["stderr"] = stderr_buf.getvalue()
print("\n" + json.dumps(result))
'''

    fd, script_path = tempfile.mkstemp(suffix=".py", dir=sandbox_dir)
    os.close(fd)
    try:
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(wrapper)

        proc = subprocess.run(
            [sys.executable, script_path],
            capture_output=True, text=True, timeout=30, cwd=sandbox_dir
        )

        # Extraire le JSON de la dernière ligne
        lines = proc.stdout.strip().split('\n')
        data = None
        for line in reversed(lines):
            line = line.strip()
            if line.startswith('{'):
                try:
                    data = json.loads(line)
                    break
                except Exception:
                    continue

        if data is None:
            return f"⚠️ Erreur d'exécution (pas de JSON retourné).\n\n**Sortie brute :**\n```\n{proc.stdout[:1500]}\n```\n**Erreur :**\n```\n{proc.stderr[:500]}\n```"

        parts = []
        if data.get("output", "").strip():
            parts.append(f"📤 **Sortie standard :**\n```\n{data['output'].strip()[:3000]}\n```")
        if data.get("stderr", "").strip():
            parts.append(f"⚠️ **Messages :**\n```\n{data['stderr'].strip()[:1000]}\n```")
        if data.get("error", "").strip():
            parts.append(f"❌ **Erreur Python :**\n```\n{data['error'].strip()[:2000]}\n```")

        files = data.get("files", [])
        if files:
            parts.append(f"📁 **Fichiers créés :** {', '.join(os.path.basename(f) for f in files)}")

        if not parts:
            parts.append("✅ Code exécuté avec succès (aucune sortie).")

        return "\n\n".join(parts)

    except subprocess.TimeoutExpired:
        return "⏱️ **ERREUR** : Le code a dépassé le temps d'exécution maximum (30 secondes). Vérifiez les boucles infinies."
    except Exception as e:
        return f"⚠️ **ERREUR lors de l'exécution** : {e}"
    finally:
        if os.path.exists(script_path):
            os.remove(script_path)


# --- MÉMOIRE VECTORIELLE (SOUVENIRS) ---
_souvenirs_collection = None

def init_souvenirs():
    """Initialise la base de données vectorielle des souvenirs."""
    global _souvenirs_collection
    if not HAS_CHROMADB:
        return False
    try:
        ef = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
        client = chromadb.PersistentClient(path="chroma_souvenirs")
        _souvenirs_collection = client.get_or_create_collection(
            name="memoire_long_terme",
            embedding_function=ef
        )
        return True
    except Exception:
        _souvenirs_collection = None
        return False


def stocker_souvenirs(facts: list[str]):
    """Stocke une liste de faits dans la mémoire vectorielle."""
    if not _souvenirs_collection or not facts:
        return
    import uuid
    ids = [f"s_{uuid.uuid4().hex[:8]}" for _ in facts]
    metas = [{"date": datetime.now().isoformat()} for _ in facts]
    try:
        _souvenirs_collection.add(documents=facts, metadatas=metas, ids=ids)
    except Exception:
        pass


def chercher_souvenirs(query: str, n: int = 3) -> list[str]:
    """Recherche les souvenirs les plus pertinents pour une question."""
    if not _souvenirs_collection:
        return []
    try:
        results = _souvenirs_collection.query(query_texts=[query], n_results=n)
        docs = results.get("documents", [[]])[0]
        return [d for d in docs if d] if docs else []
    except Exception:
        return []


def extraire_souvenirs(user_msg: str, assistant_msg: str) -> list[str]:
    """Demande au modèle d'extraire 1-3 faits clés de la conversation."""
    prompt = [
        {"role": "system", "content": "Tu es un extracteur de faits personnels. Résume cette conversation en 1 à 3 faits très courts (max 20 mots chacun) sur l'utilisateur. Exemples : 'Le chat de l'utilisateur s'appelle Rouxy', 'L'utilisateur travaille sur un projet Arduino', 'L'utilisateur préfère Python à JavaScript'. Si aucun fait personnel, réponds UNIQUEMENT 'AUCUN'. Ne réponds jamais autre chose que les faits ou AUCUN."},
        {"role": "user", "content": f"Utilisateur : {user_msg}\nAssistant : {assistant_msg}\nFaits :"}
    ]
    try:
        resp = ollama.chat(model=MODEL_TEXT, messages=prompt)
        text = resp.message.content.strip() if resp.message.content else ""
        if text.upper().startswith("AUCUN") or not text:
            return []
        return [f.strip("- • ").strip() for f in text.split("\n") if f.strip() and len(f.strip()) > 5]
    except Exception:
        return []


# --- CRÉATION DOCUMENTS ---

def creer_document(nom_fichier: str, contenu: str, format: str = "txt") -> str:
    _ensure_outputs_dir()
    safe_name = _sanitize_filename(nom_fichier)
    if not safe_name:
        safe_name = "document.txt"
    path = _get_safe_path(safe_name)
    ext = os.path.splitext(safe_name)[1].lower()

    if not format or format == "auto":
        if ext in (".txt", ".md", ".csv", ".json", ".docx", ".pdf"):
            format = ext.lstrip(".")
        else:
            format = "txt"
    format = format.lower().strip()

    try:
        if format == "docx":
            if not HAS_DOCX:
                return "Erreur : python-docx non installé."
            doc = Document()
            for line in contenu.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
            doc.save(path)

        elif format == "pdf":
            if not HAS_FPDF:
                return "Erreur : fpdf2 non installé."
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            try:
                pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
                pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
                pdf.set_font("DejaVu", size=12)
            except Exception:
                pdf.set_font("Arial", size=12)

            for line in contenu.split("\n"):
                line = line.strip()
                if not line:
                    pdf.ln(5)
                    continue
                if line.startswith("# "):
                    try:
                        pdf.set_font("DejaVu", "B", 16)
                    except Exception:
                        pdf.set_font("Arial", "B", 16)
                    pdf.cell(0, 10, line[2:], ln=True)
                    pdf.ln(2)
                    try:
                        pdf.set_font("DejaVu", size=12)
                    except Exception:
                        pdf.set_font("Arial", size=12)
                elif line.startswith("## "):
                    try:
                        pdf.set_font("DejaVu", "B", 14)
                    except Exception:
                        pdf.set_font("Arial", "B", 14)
                    pdf.cell(0, 8, line[3:], ln=True)
                    pdf.ln(2)
                    try:
                        pdf.set_font("DejaVu", size=12)
                    except Exception:
                        pdf.set_font("Arial", size=12)
                else:
                    pdf.multi_cell(0, 6, line)
                    pdf.ln(1)
            pdf.output(path)

        elif format == "csv":
            with open(path, "w", newline="", encoding="utf-8") as f:
                f.write(contenu)

        elif format == "json":
            try:
                data = json.loads(contenu)
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            except Exception:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(contenu)

        else:
            with open(path, "w", encoding="utf-8") as f:
                f.write(contenu)

        return f"Document créé avec succès : {path} ({format.upper()})"
    except Exception as e:
        return f"Erreur lors de la création du document : {e}"


AVAILABLE_FUNCTIONS = {
    "calculer": calculer,
    "heure_actuelle": heure_actuelle,
    "creer_document": creer_document,
    "rechercher_web": rechercher_web,
    "executer_python": executer_python,
    "ajouter_tache": ajouter_tache,
    "lister_taches": lister_taches,
    "marquer_fait": marquer_fait,
    "supprimer_tache": supprimer_tache,
}

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "calculer",
            "description": "Calcule une expression mathématique simple.",
            "parameters": {
                "type": "object",
                "properties": {"expression": {"type": "string", "description": "Expression mathématique"}},
                "required": ["expression"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "heure_actuelle",
            "description": "Retourne la date et l'heure actuelle.",
            "parameters": {"type": "object", "properties": {}}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "rechercher_web",
            "description": (
                "Effectue une recherche sur le web via DuckDuckGo pour obtenir des informations à jour. "
                "À utiliser uniquement quand l'utilisateur pose une question sur l'actualité, la météo, "
                "les cours boursiers, les événements récents, ou tout sujet qui pourrait avoir changé "
                "après 2024. La recherche est gratuite et ne nécessite pas de clé API."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "requete": {"type": "string", "description": "La requête de recherche (ex: 'météo Paris', 'cours Bitcoin aujourd'hui', 'résultat match France')"},
                    "n_resultats": {"type": "integer", "description": "Nombre de résultats souhaités (1-10, défaut 3)"}
                },
                "required": ["requete"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "ajouter_tache",
            "description": (
                "Ajoute une tâche ou un rappel à l'agenda de l'utilisateur. "
                "À utiliser quand l'utilisateur demande de se souvenir, de rappeler, d'ajouter à sa liste, "
                "de planifier, ou de noter une tâche à faire."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "nom": {"type": "string", "description": "Titre court de la tâche (ex: 'Appeler le dentiste', 'Acheter du lait')"},
                    "description": {"type": "string", "description": "Détails optionnels de la tâche"},
                    "date": {"type": "string", "description": "Date d'échéance au format YYYY-MM-DD (ex: '2026-06-25') ou laisser vide"},
                    "priorite": {"type": "string", "description": "basse, normale, haute, urgente. Défaut: normale"}
                },
                "required": ["nom"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "lister_taches",
            "description": (
                "Liste les tâches de l'agenda. À utiliser quand l'utilisateur demande ce qu'il doit faire, "
                "son emploi du temps, sa liste de courses, ou ses rappels."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filtre_date": {"type": "string", "description": "Filtrer par date (YYYY-MM-DD) ou laisser vide pour toutes"},
                    "filtre_statut": {"type": "string", "description": "'afaire' pour non-faites, 'fait' pour terminées, ou vide pour toutes"},
                    "limite": {"type": "integer", "description": "Nombre maximum de tâches à afficher (1-50). Défaut: 20"}
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "marquer_fait",
            "description": (
                "Marque une tâche comme terminée. À utiliser quand l'utilisateur dit qu'il a fait quelque chose, "
                "ou qu'il veut cocher une tâche."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "tache_id": {"type": "string", "description": "ID de la tâche (ex: 't_1234_5678') ou nom partiel de la tâche"}
                },
                "required": ["tache_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "supprimer_tache",
            "description": (
                "Supprime une tâche de l'agenda. À utiliser quand l'utilisateur veut annuler un rappel "
                "ou supprimer une tâche de sa liste."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "tache_id": {"type": "string", "description": "ID de la tâche ou nom partiel"}
                },
                "required": ["tache_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "executer_python",
            "description": (
                "Exécute du code Python dans un bac à sable sécurisé. "
                "Utilise cet outil pour les calculs complexes, les graphiques (matplotlib), "
                "les analyses statistiques (numpy/pandas), ou les transformations de données. "
                "Le code ne peut pas accéder au système, au réseau, ou à vos fichiers personnels. "
                "Les graphiques matplotlib sont automatiquement sauvegardés en PNG dans outputs/sandbox/."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {
                        "type": "string",
                        "description": "Code Python complet à exécuter. N'incluez pas d'imports dangereux (os, subprocess, socket, etc.). Matplotlib, numpy, pandas sont disponibles."
                    },
                    "nom_fichier_sortie": {
                        "type": "string",
                        "description": "Nom optionnel pour les graphiques générés (ex: 'histogramme.png'). Défaut : figure_1.png"
                    }
                },
                "required": ["code"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "creer_document",
            "description": (
                "Crée un fichier document (txt, md, csv, json, docx, pdf) dans le dossier outputs/. "
                "À utiliser quand l'utilisateur demande de sauvegarder, exporter, générer un fichier, "
                "créer un rapport, une lettre, un tableau de données, etc."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "nom_fichier": {"type": "string", "description": "Nom avec extension (ex: rapport.txt, données.csv, lettre.docx, resume.pdf)"},
                    "contenu": {"type": "string", "description": "Contenu textuel complet. Pour CSV, utiliser des virgules et retours à la ligne."},
                    "format": {"type": "string", "description": "Format : txt, md, csv, json, docx, pdf. Déduit de l'extension si non fourni."}
                },
                "required": ["nom_fichier", "contenu"]
            }
        }
    }
]

if not HAS_DUCKDUCKGO:
    # Retirer l'outil recherche web si la bibliothèque n'est pas installée
    TOOLS = [t for t in TOOLS if t.get("function", {}).get("name") != "rechercher_web"]

if HAS_CHROMADB:
    TOOLS.append({
        "type": "function",
        "function": {
            "name": "rechercher_documents",
            "description": "Recherche des passages dans les documents locaux (PDF, Word, TXT).",
            "parameters": {
                "type": "object",
                "properties": {
                    "requete": {"type": "string", "description": "Sujet à chercher"},
                    "n_resultats": {"type": "integer", "description": "Nombre de passages (1-10)"}
                },
                "required": ["requete"]
            }
        }
    })
    AVAILABLE_FUNCTIONS["rechercher_documents"] = rechercher_documents


# ==============================================================================
# 4. PRÉPARATION DU SYSTEM PROMPT MULTI-LANGUE
# ==============================================================================
def preparer_system_prompt(persona_nom: str) -> str:
    target_language = LANG_NAMES.get(st.session_state.lang, "Français (French)")
    
    # Injection dynamique stricte pour forcer le LLM à basculer
    full_prompt = (
        f"--- IMPORTANT STRICT DIRECTIVE ---\n"
        f"You must strictly analyze, think, and communicate your final response to the user entirely in this language: {target_language}.\n"
        f"Do not use any other language for your final response."
    )
    return full_prompt



# ========================= 5. MÉMOIRE PERSISTANTE =========================

def load_memory():
    if not os.path.exists(MEMORY_FILE):
        return []
    try:
        with open(MEMORY_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data if isinstance(data, list) else []
    except Exception:
        return []


def save_memory(messages):
    try:
        with open(MEMORY_FILE, "w", encoding="utf-8") as f:
            json.dump(messages, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.sidebar.warning(f"Erreur sauvegarde mémoire : {e}")


def summarize_text_history(messages):
    if len(messages) <= MAX_TEXT_HISTORY:
        return messages
    recent = messages[-6:]
    old = messages[:-6]
    summary_prompt = [
        {"role": "system", "content": "Résume en quelques phrases les informations clés (prénom, préférences, sujets, décisions). Sois très concis."},
        {"role": "user", "content": "Historique à résumer :\n\n" + "\n".join(
            f"{m.get('role','?')}: {m.get('content','')}" for m in old if m.get("content")
        )}
    ]
    try:
        resp = ollama.chat(model=MODEL_TEXT, messages=summary_prompt)
        summary_text = resp.message.content or "(pas de résumé)"
    except Exception as e:
        summary_text = f"(erreur résumé: {e})"
    new_msgs = [{"role": "system", "content": f"Résumé des conversations précédentes : {summary_text}"}]
    new_msgs.extend(recent)
    return new_msgs


# ========================= 6. ENCODAGE IMAGE =========================

def encode_image(uploaded_file):
    raw = uploaded_file.getvalue()
    if HAS_PIL:
        try:
            img = Image.open(BytesIO(raw))
            w, h = img.size
            if w > MAX_IMAGE_SIZE or h > MAX_IMAGE_SIZE:
                ratio = min(MAX_IMAGE_SIZE / w, MAX_IMAGE_SIZE / h)
                new_w, new_h = int(w * ratio), int(h * ratio)
                img = img.resize((new_w, new_h), Image.LANCZOS)
                st.sidebar.info(f"Image redimensionnée : {w}x{h} → {new_w}x{new_h}")
            if img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            buf = BytesIO()
            img.save(buf, format="PNG")
            raw = buf.getvalue()
        except Exception as e:
            st.sidebar.warning(f"Erreur redimensionnement : {e}")
    else:
        st.sidebar.info("Pillow non installé. Image envoyée telle quelle.")
    return base64.b64encode(raw).decode("utf-8")


# ========================= 7. AGENT TEXTE =========================

def _parse_args(args):
    if isinstance(args, dict):
        return args
    if isinstance(args, str):
        try:
            return json.loads(args)
        except Exception:
            return {}
    return {}


def run_text_agent(history, user_msg, system_prompt: str = None):
    """
    Gère un tour de conversation texte avec tool calling.
    Injecte le system_prompt au début de l'historique si fourni.
    """
    # Injecter le system prompt si présent et non déjà en place
    if system_prompt:
        has_system = any(m.get("role") == "system" for m in history)
        if not has_system:
            history.insert(0, {"role": "system", "content": system_prompt})
        else:
            # Mettre à jour le system prompt existant
            for i, m in enumerate(history):
                if m.get("role") == "system":
                    history[i]["content"] = system_prompt
                    break

    history.append({"role": "user", "content": user_msg})

    # Chercher et injecter des souvenirs pertinents (mémoire vectorielle)
    souvenirs = chercher_souvenirs(user_msg)
    if souvenirs:
        history.append({"role": "system", "content": f"Souvenirs pertinents de conversations précédentes : {' | '.join(souvenirs)}"})

    response = ollama.chat(model=MODEL_TEXT, messages=history, tools=TOOLS)

    while response.message.tool_calls:
        assistant_msg = {
            "role": "assistant",
            "content": response.message.content or "",
            "tool_calls": [
                {
                    "type": "function",
                    "function": {
                        "name": tc.function.name,
                        "arguments": tc.function.arguments
                    }
                }
                for tc in response.message.tool_calls
            ]
        }
        history.append(assistant_msg)

        for tool_call in response.message.tool_calls:
            func_name = tool_call.function.name
            func_args = _parse_args(tool_call.function.arguments)
            func = AVAILABLE_FUNCTIONS.get(func_name)

            if func is None:
                result = f"Erreur : outil '{func_name}' introuvable."
            else:
                try:
                    result = func(**func_args)
                except Exception as e:
                    result = f"Erreur : {e}"

            history.append({"role": "tool", "name": func_name, "content": str(result)})

        response = ollama.chat(model=MODEL_TEXT, messages=history, tools=TOOLS)

    final = response.message.content or "(pas de réponse)"
    history.append({"role": "assistant", "content": final})
    return final


# ========================= 8. AGENT VISION =========================

def extract_data_from_image(image_b64: str) -> str:
    """
    Appelle LLaVA pour extraire les données brutes d'une image
    (tableaux, chiffres, graphiques, texte structuré).
    Retourne un texte brut avec les données extraites.
    """
    try:
        response = ollama.chat(
            model=MODEL_VISION,
            messages=[{
                "role": "user",
                "content": (
                    "Extrais TOUTES les données textuelles et numériques de cette image. "
                    "Présente-les sous forme brute structurée (CSV, liste, ou tableau texte). "
                    "Ne fais AUCUNE analyse, ne donne AUCUNE opinion. "
                    "Juste les données brutes exactement comme elles apparaissent."
                ),
                "images": [image_b64]
            }]
        )
        return response.message.content or "(aucune donnée extraite)"
    except Exception as e:
        return f"Erreur extraction vision : {e}"


def run_vision_agent(vision_history, user_msg, image_b64):
    msgs = vision_history[-(MAX_VISION_HISTORY - 2):]
    msgs.append({"role": "user", "content": user_msg, "images": [image_b64]})
    try:
        response = ollama.chat(model=MODEL_VISION, messages=msgs)
        reply = response.message.content or "(pas de réponse)"
    except Exception as e:
        reply = f"[ERREUR Vision] {e}"
    vision_history.append({"role": "user", "content": user_msg})
    vision_history.append({"role": "assistant", "content": reply})
    while len(vision_history) > MAX_VISION_HISTORY:
        vision_history.pop(0)
    return reply


# ========================= 9. INTERFACE STREAMLIT =========================

def main():
    st.set_page_config(page_title="Agent Ultime IA", page_icon="🌟", layout="centered")
    st.title("🌟🤖👁️🎙️📝 Agent Ultime IA")
    st.caption("Texte · Vision · Voix · Micro · Documents · Mémoire | 100% local")

    # --- Initialisation singletons ---
    if "tts" not in st.session_state:
        st.session_state.tts = TTSManager(use_piper=USE_PIPER)
    if "stt" not in st.session_state:
        st.session_state.stt = STTManager(use_whisper=USE_FASTER_WHISPER)

    tts = st.session_state.tts
    stt = st.session_state.stt

    # ==================== SIDEBAR ====================
    with st.sidebar:
        st.header("⚙️ Configuration")
        st.markdown(f"**🧠 Texte** : `{MODEL_TEXT}`")
        st.markdown(f"**👁️ Vision** : `{MODEL_VISION}`")
        st.divider()

        # Upload image
        uploaded_file = st.file_uploader(
            "📁 Uploader une image",
            type=["png", "jpg", "jpeg", "gif", "bmp", "webp"]
        )
        if uploaded_file is not None:
            st.image(uploaded_file, caption="Image prête", use_container_width=True)
            st.success("Mode vision activé")

        st.divider()

        # SÉLECTEUR DE LANGUE EN DIRECT
        lang_display = {"fr": "🇫🇷 Français", "en": "🇬🇧 English", "es": "🇪🇸 Español", "de": "🇩🇪 Deutsch"}
        selected_lang_display = st.selectbox(
            t("select_lang"),
            options=list(lang_display.values()),
            index=list(lang_display.keys()).index(st.session_state.lang)
        )
        # Vérification et mise à jour instantanée si la langue change
        for code, disp in lang_display.items():
            if disp == selected_lang_display and st.session_state.lang != code:
                st.session_state.lang = code
                st.rerun()  # Recharge la page instantanément avec la nouvelle langue

        st.divider()
        
        # Personnalité
        personnalites = load_personnalites()
        noms_personnalites = [p["nom"] for p in personnalites]
        if "personnalite_auto" not in st.session_state:
            st.session_state.personnalite_auto = True
        if "personnalite_idx" not in st.session_state:
            st.session_state.personnalite_idx = 0
        if "personnalite_detected_idx" not in st.session_state:
            st.session_state.personnalite_detected_idx = 0

        auto_detect = st.checkbox("🤖 Détection auto des personnalités", value=st.session_state.personnalite_auto)

        if auto_detect:
            st.session_state.personnalite_auto = True
            detected_idx = st.session_state.get("personnalite_detected_idx", 0)
            detected_nom = noms_personnalites[detected_idx] if 0 <= detected_idx < len(noms_personnalites) else noms_personnalites[0]
            st.success(f"🎭 Actuel : {detected_nom}")
            st.caption("L'agent change de ton selon vos questions.")
            personnalite_idx = detected_idx
            personnalite_choisie = personnalites[personnalite_idx]
        else:
            st.session_state.personnalite_auto = False
            personnalite_idx = st.selectbox(
                "🎭 Personnalité (manuel)",
                range(len(noms_personnalites)),
                format_func=lambda i: noms_personnalites[i],
                index=st.session_state.personnalite_idx
            )
            st.caption(personnalites[personnalite_idx].get("description", ""))
            if personnalite_idx != st.session_state.personnalite_idx:
                st.session_state.personnalite_idx = personnalite_idx
                for key in ["messages", "internal_history"]:
                    if key in st.session_state:
                        del st.session_state[key]
                st.info(f"🎭 Personnalité changée : {personnalites[personnalite_idx]['nom']}. Historique texte réinitialisé.")
            personnalite_choisie = personnalites[personnalite_idx]

        st.divider()

        # Indicateur connexion Internet
        is_online = check_internet()
        if is_online:
            st.success("🌐 Connexion Internet : OK")
        else:
            st.warning("🌐 Connexion Internet : OFFLINE")

        st.divider()

        # Options
        use_memory = st.checkbox("🧠 Mémoire persistante", value=False)
        use_rag = st.checkbox("📄 Recherche documents", value=False)
        use_web = st.checkbox("🌐 Activer recherche web (DuckDuckGo)", value=False, disabled=not is_online)
        enable_tts = st.checkbox("🔊 Lire les réponses à voix haute", value=tts.ok)
        enable_stt = st.checkbox("🎤 Activer le micro (bouton ci-dessous)", value=False)

        if use_rag:
            if init_rag():
                st.success("Base documents OK.")
            else:
                st.warning("Base introuvable. Lancez agent_documents.py d'abord.")

        # Mémoire vectorielle (souvenirs) - initialisation silencieuse
        souvenirs_ready = init_souvenirs()
        if souvenirs_ready:
            st.success("🧠 Mémoire vectorielle prête.")
        else:
            st.info("🧠 Mémoire vectorielle : installez `pip install sentence-transformers` pour activer les souvenirs long terme.")

        st.divider()

        # Bouton micro
        if enable_stt and stt.ok:
            if st.button("🎙️ Parler maintenant", use_container_width=True):
                st.session_state.trigger_micro = True
            st.caption("Cliquez, puis parlez. Attendez le traitement.")
        elif enable_stt:
            st.warning("Micro non disponible. Installez pyaudio + SpeechRecognition (+ faster-whisper pour offline).")

        st.divider()

        # Nouvelle conversation
        if st.button("🗑️ Nouvelle conversation", use_container_width=True):
            for key in ["messages", "internal_history", "vision_history", "current_image_b64", "current_image_name"]:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()

        st.divider()

        # Documents créés
        st.subheader("📁 Documents créés")
        if os.path.exists(OUTPUTS_DIR):
            files = sorted(os.listdir(OUTPUTS_DIR))
            if files:
                for f in files:
                    fpath = os.path.join(OUTPUTS_DIR, f)
                    fsize = os.path.getsize(fpath)
                    st.caption(f"📄 {f} ({fsize:,} octets)")
            else:
                st.caption("Aucun document pour l'instant.")
        else:
            st.caption("Dossier outputs/ vide.")

    # ==================== INIT ÉTAT ====================
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "internal_history" not in st.session_state:
        st.session_state.internal_history = []
    if "vision_history" not in st.session_state:
        st.session_state.vision_history = []
    if "current_image_b64" not in st.session_state:
        st.session_state.current_image_b64 = None
    if "current_image_name" not in st.session_state:
        st.session_state.current_image_name = None
    if "trigger_micro" not in st.session_state:
        st.session_state.trigger_micro = False
    if "personnalite_idx" not in st.session_state:
        st.session_state.personnalite_idx = 0

    # Charger mémoire
    if use_memory and not st.session_state.messages:
        loaded = load_memory()
        if loaded:
            st.session_state.internal_history = loaded
            for i in range(0, len(loaded), 2):
                if i < len(loaded) and loaded[i].get("role") == "user":
                    st.session_state.messages.append({"role": "user", "content": loaded[i].get("content", "")})
                if i + 1 < len(loaded) and loaded[i+1].get("role") == "assistant":
                    st.session_state.messages.append({"role": "assistant", "content": loaded[i+1].get("content", "")})
            st.sidebar.info(f"Mémoire chargée : {len(loaded)} messages")

    # ==================== AFFICHAGE HISTORIQUE ====================
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            if msg.get("image_b64"):
                st.image(BytesIO(base64.b64decode(msg["image_b64"])), use_container_width=True)
            if msg.get("content"):
                st.markdown(msg["content"])

    # ==================== TRAITEMENT MICRO ====================
        if st.session_state.trigger_micro and enable_stt and stt.ok:
            st.session_state.trigger_micro = False
            with st.spinner("🎤 J'écoute..."):
                heard = stt.listen()
            if heard:
                st.info(f"👤 Vous avez dit : *{heard}*")
                # Détection auto de la personnalité avant traitement
                if auto_detect:
                    st.session_state.personnalite_detected_idx = detecter_personnalite_auto(heard, personnalites)
                    personnalite_choisie = personnalites[st.session_state.personnalite_detected_idx]
                # Traiter comme input utilisateur
                process_input(heard, uploaded_file, use_memory, enable_tts, tts, system_prompt=personnalite_choisie.get("system_prompt"), personnalite_nom=personnalite_choisie.get("nom") if auto_detect else None)
            else:
                st.warning("Je n'ai pas compris. Veuillez réessayer ou taper votre message.")

    # ==================== INPUT TEXTE ====================
    user_input = st.chat_input("Posez votre question, décrivez une image, ou demandez un document...")

    if user_input:
        # Détection auto de la personnalité avant traitement
        if auto_detect:
            st.session_state.personnalite_detected_idx = detecter_personnalite_auto(user_input, personnalites)
            personnalite_choisie = personnalites[st.session_state.personnalite_detected_idx]
        process_input(user_input, uploaded_file, use_memory, enable_tts, tts, system_prompt=personnalite_choisie.get("system_prompt"), personnalite_nom=personnalite_choisie.get("nom") if auto_detect else None)


def process_input(user_input: str, uploaded_file, use_memory: bool, enable_tts: bool, tts: TTSManager, system_prompt: str = None, personnalite_nom: str = None):
    """Fonction centrale qui traite l'input utilisateur (texte ou micro)."""

    if uploaded_file is not None:
        b64 = encode_image(uploaded_file)
        st.session_state.current_image_b64 = b64
        st.session_state.current_image_name = uploaded_file.name

    has_image = st.session_state.current_image_b64 is not None
    is_vision = has_image and user_input

    user_display_text = user_input if user_input else "Analyse cette image."

    with st.chat_message("user"):
        if has_image:
            st.image(BytesIO(base64.b64decode(st.session_state.current_image_b64)), use_container_width=True)
        st.markdown(user_display_text)

    st.session_state.messages.append({
        "role": "user",
        "content": user_display_text,
        "image_b64": st.session_state.current_image_b64 if has_image else None
    })

    with st.chat_message("assistant"):
        with st.spinner("Analyse en cours..."):
            try:
                if is_vision:
                    # Détection pipeline Vision → Analyse de données (Python)
                    mots_pipeline = [
                        "graphique", "analyse", "analyser", "données", "csv", "calcul",
                        "moyenne", "médiane", "statistique", "plot", "chart", "histogramme",
                        "barres", "courbe", "total", "somme", "dataframe", "pandas", "numpy",
                        "tableau", "excel", "chiffre", "nombre", "montant", "score",
                        "pourcentage", "taux", "comparer", "évolution", "croissance",
                        "décroissance", "minimum", "maximum", "min", "max", "tri", "rang", "classement"
                    ]
                    declenche_pipeline = user_input and any(m in user_input.lower() for m in mots_pipeline)

                    if declenche_pipeline:
                        # Étape 1 : Extraction des données brutes avec LLaVA
                        st.markdown("🔍 *Extraction des données de l'image...*")
                        extraction = extract_data_from_image(st.session_state.current_image_b64)

                        # Étape 2 : Analyse avec Qwen + bac à sable Python
                        pipeline_prompt = (
                            f"Voici les données brutes extraites de l'image :\n"
                            f"---\n{extraction}\n---\n\n"
                            f"{user_display_text}"
                        )
                        if len(st.session_state.internal_history) > MAX_TEXT_HISTORY:
                            st.session_state.internal_history = summarize_text_history(
                                st.session_state.internal_history
                            )
                            st.sidebar.info("Historique texte résumé.")
                        reply = run_text_agent(
                            st.session_state.internal_history,
                            pipeline_prompt,
                            system_prompt=system_prompt
                        )
                    else:
                        reply = run_vision_agent(
                            st.session_state.vision_history,
                            user_display_text,
                            st.session_state.current_image_b64
                        )
                else:
                    if len(st.session_state.internal_history) > MAX_TEXT_HISTORY:
                        st.session_state.internal_history = summarize_text_history(
                            st.session_state.internal_history
                        )
                        st.sidebar.info("Historique texte résumé.")
                    reply = run_text_agent(
                        st.session_state.internal_history,
                        user_display_text,
                        system_prompt=system_prompt
                    )
            except Exception as e:
                reply = f"❌ Erreur : {e}"

        st.markdown(reply)

        if personnalite_nom:
            st.caption(f"🎭 {personnalite_nom}")

        # Afficher les graphiques/images générés par le bac à sable
        sandbox_dir = os.path.join(OUTPUTS_DIR, "sandbox")
        if os.path.exists(sandbox_dir):
            img_files = [f for f in os.listdir(sandbox_dir) if f.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp"))]
            if img_files:
                st.markdown("📊 **Graphiques générés :**")
                for img_name in sorted(img_files):
                    img_path = os.path.join(sandbox_dir, img_name)
                    st.image(img_path, caption=img_name, use_container_width=True)

        if "Document créé avec succès" in reply or "créé avec succès" in reply:
            st.success("📄 Document généré ! Vérifiez la sidebar.")
            st.rerun()  # Rafraîchir la liste des documents

    st.session_state.messages.append({"role": "assistant", "content": reply})

    if use_memory:
        save_memory(st.session_state.internal_history)

    # Extraction et stockage automatique des souvenirs (mémoire vectorielle)
    if not is_vision and reply:
        facts = extraire_souvenirs(user_display_text, reply)
        if facts:
            stocker_souvenirs(facts)
            st.sidebar.info(f"🧠 {len(facts)} souvenir(s) stocké(s).")

    if enable_tts and tts.ok:
        tts.speak(reply)


if __name__ == "__main__":
    main()
