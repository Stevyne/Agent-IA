"""
🤖👁️📝 Agent Hybride Web V2 — Texte + Vision + Outils + Mémoire + Export Documents

Interface Streamlit qui route automatiquement :
  - Texte seul      → qwen2.5:3b (outils : calcul, heure, RAG, créer_document)
  - Image présente  → llava-phi3 (vision)

Capacité d'export : TXT, MD, CSV, JSON, DOCX, PDF dans le dossier outputs/.

INSTALLATION :
  pip install streamlit pillow ollama python-docx fpdf2
  ollama pull qwen2.5:3b
  ollama pull llava-phi3
"""

import streamlit as st
import ollama
import os
import sys
import json
import csv
import base64
import re
from datetime import datetime
from io import BytesIO

# ========================= CONFIGURATION =========================
MODEL_TEXT = "qwen2.5:3b"
MODEL_VISION = "llava-phi3"
MEMORY_FILE = "memory_hybride.json"
OUTPUTS_DIR = "outputs"
MAX_TEXT_HISTORY = 20
MAX_VISION_HISTORY = 10
MAX_IMAGE_SIZE = 1024
RAG_DB_PATH = "chroma_db_docs"
RAG_COLLECTION = "bibliotheque"
# ================================================================


# --- Imports optionnels ---
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import chromadb
    HAS_CHROMADB = True
except ImportError:
    HAS_CHROMADB = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False


# ------------------------------------------------------------------
# 1. SÉCURITÉ : BAC À SABLE outputs/
# ------------------------------------------------------------------

def _sanitize_filename(name: str) -> str:
    """
    Nettoie un nom de fichier pour éviter les attaques de chemin.
    Supprime les .., les slashs absolus, et les caractères dangereux.
    """
    if not name:
        return "document.txt"
    # Remplacer les caractères interdits
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    # Supprimer les .. et les chemins absolus
    name = name.replace("..", "_")
    name = name.lstrip("/\\")
    return name


def _ensure_outputs_dir():
    """Crée le dossier outputs/ s'il n'existe pas."""
    if not os.path.exists(OUTPUTS_DIR):
        os.makedirs(OUTPUTS_DIR)


def _get_safe_path(name: str) -> str:
    """Retourne un chemin sûr dans outputs/ pour le nom de fichier donné."""
    _ensure_outputs_dir()
    safe = _sanitize_filename(name)
    return os.path.join(OUTPUTS_DIR, safe)


# ------------------------------------------------------------------
# 2. OUTILS TEXTE
# ------------------------------------------------------------------

def calculer(expression: str) -> str:
    if not re.match(r'^[0-9+\-*/.() ]+$', expression):
        return "Erreur : caractères non autorisés."
    try:
        return str(eval(expression))
    except Exception as e:
        return f"Erreur de calcul : {e}"


def heure_actuelle() -> str:
    return datetime.now().strftime("%A %d %B %Y, %H:%M:%S")


# --- RAG ---
_rag_collection = None

def init_rag():
    global _rag_collection
    if not HAS_CHROMADB:
        return False
    try:
        client = chromadb.PersistentClient(path=RAG_DB_PATH)
        _rag_collection = client.get_collection(name=RAG_COLLECTION)
        return True
    except Exception:
        _rag_collection = None
        return False


def rechercher_documents(requete: str, n_resultats: int = 5) -> str:
    if _rag_collection is None:
        return "Base de documents non initialisée."
    try:
        n_resultats = int(n_resultats)
        if n_resultats < 1:
            n_resultats = 1
        if n_resultats > 10:
            n_resultats = 10
    except Exception:
        n_resultats = 5

    try:
        results = _rag_collection.query(query_texts=[requete], n_results=n_resultats)
    except Exception as e:
        return f"Erreur recherche : {e}"

    if not results or not results["documents"] or not results["documents"][0]:
        return "Aucun passage pertinent trouvé."

    passages = []
    for i, doc in enumerate(results["documents"][0]):
        meta = results["metadatas"][0][i] if results.get("metadatas") else {}
        source = meta.get("source", "inconnu")
        doc_type = meta.get("type", "?")
        page = meta.get("page", "?")

        if doc_type == "pdf":
            loc = f"page {page}"
        elif doc_type == "docx":
            loc = f"paragraphe {page}"
        elif doc_type == "txt":
            loc = f"bloc {page}"
        else:
            loc = f"position {page}"

        passages.append(f"---\nSource: {source} ({loc})\n{doc}")

    return "\n".join(passages)


# --- CRÉATION DE DOCUMENTS ---

def creer_document(nom_fichier: str, contenu: str, format: str = "txt") -> str:
    """
    Crée un document dans le dossier outputs/.
    Formats supportés : txt, md, csv, json, docx, pdf.
    """
    _ensure_outputs_dir()
    safe_name = _sanitize_filename(nom_fichier)
    if not safe_name:
        safe_name = "document.txt"

    path = _get_safe_path(safe_name)
    ext = os.path.splitext(safe_name)[1].lower()

    # Détecter le format depuis l'extension si non fourni
    if not format or format == "auto":
        if ext in (".txt", ".md", ".csv", ".json", ".docx", ".pdf"):
            format = ext.lstrip(".")
        else:
            format = "txt"

    format = format.lower().strip()

    try:
        if format == "docx":
            if not HAS_DOCX:
                return "Erreur : python-docx non installé. Document non créé."
            doc = Document()
            # Traiter le contenu ligne par ligne pour détecter les titres
            for line in contenu.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
            doc.save(path)

        elif format == "pdf":
            if not HAS_FPDF:
                return "Erreur : fpdf2 non installé. Document non créé."
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            # Essayer de charger une police UTF-8 supportant l'accès aux caractères français
            # fpdf2 embarque DejaVuSans qui supporte les accents
            try:
                pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
                pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
                pdf.set_font("DejaVu", size=12)
            except Exception:
                # Fallback : police standard (accents limités)
                pdf.set_font("Arial", size=12)

            for line in contenu.split("\n"):
                line = line.strip()
                if not line:
                    pdf.ln(5)
                    continue
                if line.startswith("# "):
                    try:
                        pdf.set_font("DejaVu", "B", 16)
                    except Exception:
                        pdf.set_font("Arial", "B", 16)
                    pdf.cell(0, 10, line[2:], ln=True)
                    pdf.ln(2)
                    try:
                        pdf.set_font("DejaVu", size=12)
                    except Exception:
                        pdf.set_font("Arial", size=12)
                elif line.startswith("## "):
                    try:
                        pdf.set_font("DejaVu", "B", 14)
                    except Exception:
                        pdf.set_font("Arial", "B", 14)
                    pdf.cell(0, 8, line[3:], ln=True)
                    pdf.ln(2)
                    try:
                        pdf.set_font("DejaVu", size=12)
                    except Exception:
                        pdf.set_font("Arial", size=12)
                else:
                    pdf.multi_cell(0, 6, line)
                    pdf.ln(1)
            pdf.output(path)

        elif format == "csv":
            # Le contenu est interprété comme du texte CSV brut
            with open(path, "w", newline="", encoding="utf-8") as f:
                f.write(contenu)

        elif format == "json":
            # Essayer de parser le contenu comme JSON, sinon le sauvegarder tel quel
            try:
                data = json.loads(contenu)
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            except Exception:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(contenu)

        else:  # txt, md, ou inconnu
            with open(path, "w", encoding="utf-8") as f:
                f.write(contenu)

        return f"Document créé avec succès : {path} ({format.upper()})"

    except Exception as e:
        return f"Erreur lors de la création du document : {e}"


AVAILABLE_FUNCTIONS = {
    "calculer": calculer,
    "heure_actuelle": heure_actuelle,
    "creer_document": creer_document,
}

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "calculer",
            "description": "Calcule une expression mathématique simple.",
            "parameters": {
                "type": "object",
                "properties": {"expression": {"type": "string", "description": "Expression mathématique"}},
                "required": ["expression"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "heure_actuelle",
            "description": "Retourne la date et l'heure actuelle.",
            "parameters": {"type": "object", "properties": {}}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "creer_document",
            "description": (
                "Crée un fichier document (texte, markdown, CSV, JSON, Word ou PDF) "
                "dans le dossier de sortie. À utiliser quand l'utilisateur demande de sauvegarder, "
                "exporter, générer un fichier, créer un rapport, une lettre, un tableau de données, etc."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "nom_fichier": {
                        "type": "string",
                        "description": "Nom du fichier avec extension (ex: rapport.txt, données.csv, lettre.docx, resume.pdf)"
                    },
                    "contenu": {
                        "type": "string",
                        "description": "Contenu textuel complet du document. Pour les tableaux CSV, utiliser des virgules et retours à la ligne."
                    },
                    "format": {
                        "type": "string",
                        "description": "Format du fichier : txt, md, csv, json, docx, pdf. Déduit de l'extension si non fourni."
                    }
                },
                "required": ["nom_fichier", "contenu"]
            }
        }
    }
]

if HAS_CHROMADB:
    TOOLS.append({
        "type": "function",
        "function": {
            "name": "rechercher_documents",
            "description": (
                "Recherche des passages dans les documents locaux (PDF, Word, TXT). "
                "À utiliser si la question concerne un document, un contrat, un rapport, etc."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "requete": {"type": "string", "description": "Sujet à chercher"},
                    "n_resultats": {"type": "integer", "description": "Nombre de passages (1-10)"}
                },
                "required": ["requete"]
            }
        }
    })
    AVAILABLE_FUNCTIONS["rechercher_documents"] = rechercher_documents


# ------------------------------------------------------------------
# 3. MÉMOIRE PERSISTANTE
# ------------------------------------------------------------------

def load_memory():
    if not os.path.exists(MEMORY_FILE):
        return []
    try:
        with open(MEMORY_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data if isinstance(data, list) else []
    except Exception:
        return []


def save_memory(messages):
    try:
        with open(MEMORY_FILE, "w", encoding="utf-8") as f:
            json.dump(messages, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.sidebar.warning(f"Erreur sauvegarde mémoire : {e}")


def summarize_text_history(messages):
    if len(messages) <= MAX_TEXT_HISTORY:
        return messages
    recent = messages[-6:]
    old = messages[:-6]
    summary_prompt = [
        {"role": "system", "content": "Résume en quelques phrases les informations clés (prénom, préférences, sujets, décisions). Sois très concis."},
        {"role": "user", "content": "Historique à résumer :\n\n" + "\n".join(
            f"{m.get('role','?')}: {m.get('content','')}" for m in old if m.get("content")
        )}
    ]
    try:
        resp = ollama.chat(model=MODEL_TEXT, messages=summary_prompt)
        summary_text = resp.message.content or "(pas de résumé)"
    except Exception as e:
        summary_text = f"(erreur résumé: {e})"
    new_msgs = [{"role": "system", "content": f"Résumé des conversations précédentes : {summary_text}"}]
    new_msgs.extend(recent)
    return new_msgs


# ------------------------------------------------------------------
# 4. ENCODAGE IMAGE
# ------------------------------------------------------------------

def encode_image(uploaded_file):
    raw = uploaded_file.getvalue()
    if HAS_PIL:
        try:
            img = Image.open(BytesIO(raw))
            w, h = img.size
            if w > MAX_IMAGE_SIZE or h > MAX_IMAGE_SIZE:
                ratio = min(MAX_IMAGE_SIZE / w, MAX_IMAGE_SIZE / h)
                new_w, new_h = int(w * ratio), int(h * ratio)
                img = img.resize((new_w, new_h), Image.LANCZOS)
                st.sidebar.info(f"Image redimensionnée : {w}x{h} → {new_w}x{new_h}")
            if img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            buf = BytesIO()
            img.save(buf, format="PNG")
            raw = buf.getvalue()
        except Exception as e:
            st.sidebar.warning(f"Erreur redimensionnement : {e}")
    else:
        st.sidebar.info("Pillow non installé. Image envoyée telle quelle.")
    return base64.b64encode(raw).decode("utf-8")


# ------------------------------------------------------------------
# 5. AGENT TEXTE
# ------------------------------------------------------------------

def _parse_args(args):
    if isinstance(args, dict):
        return args
    if isinstance(args, str):
        try:
            return json.loads(args)
        except Exception:
            return {}
    return {}


def run_text_agent(history, user_msg):
    history.append({"role": "user", "content": user_msg})
    response = ollama.chat(model=MODEL_TEXT, messages=history, tools=TOOLS)

    while response.message.tool_calls:
        assistant_msg = {
            "role": "assistant",
            "content": response.message.content or "",
            "tool_calls": [
                {
                    "type": "function",
                    "function": {
                        "name": tc.function.name,
                        "arguments": tc.function.arguments
                    }
                }
                for tc in response.message.tool_calls
            ]
        }
        history.append(assistant_msg)

        for tool_call in response.message.tool_calls:
            func_name = tool_call.function.name
            func_args = _parse_args(tool_call.function.arguments)
            func = AVAILABLE_FUNCTIONS.get(func_name)

            if func is None:
                result = f"Erreur : outil '{func_name}' introuvable."
            else:
                try:
                    result = func(**func_args)
                except Exception as e:
                    result = f"Erreur : {e}"

            history.append({"role": "tool", "name": func_name, "content": str(result)})

        response = ollama.chat(model=MODEL_TEXT, messages=history, tools=TOOLS)

    final = response.message.content or "(pas de réponse)"
    history.append({"role": "assistant", "content": final})
    return final


# ------------------------------------------------------------------
# 6. AGENT VISION
# ------------------------------------------------------------------

def run_vision_agent(vision_history, user_msg, image_b64):
    msgs = vision_history[-(MAX_VISION_HISTORY - 2):]
    msgs.append({"role": "user", "content": user_msg, "images": [image_b64]})
    try:
        response = ollama.chat(model=MODEL_VISION, messages=msgs)
        reply = response.message.content or "(pas de réponse)"
    except Exception as e:
        reply = f"[ERREUR Vision] {e}"
    vision_history.append({"role": "user", "content": user_msg})
    vision_history.append({"role": "assistant", "content": reply})
    while len(vision_history) > MAX_VISION_HISTORY:
        vision_history.pop(0)
    return reply


# ------------------------------------------------------------------
# 7. INTERFACE STREAMLIT
# ------------------------------------------------------------------

def main():
    st.set_page_config(page_title="Agent Hybride IA V2", page_icon="🤖", layout="centered")
    st.title("🤖👁️📝 Agent Hybride IA V2")
    st.caption("Texte · Vision · Outils · Mémoire · Export Documents | 100% local")

    # ==================== SIDEBAR ====================
    with st.sidebar:
        st.header("⚙️ Configuration")
        st.markdown(f"**🧠 Texte** : `{MODEL_TEXT}`")
        st.markdown(f"**👁️ Vision** : `{MODEL_VISION}`")
        st.divider()

        # Upload image
        uploaded_file = st.file_uploader(
            "📁 Uploader une image",
            type=["png", "jpg", "jpeg", "gif", "bmp", "webp"]
        )
        if uploaded_file is not None:
            st.image(uploaded_file, caption="Image prête", use_container_width=True)
            st.success("Mode vision activé")

        st.divider()

        # Checkboxes
        use_memory = st.checkbox("🧠 Mémoire persistante", value=False)
        use_rag = st.checkbox("📄 Recherche documents", value=False)

        if use_rag:
            if init_rag():
                st.success("Base documents OK.")
            else:
                st.warning("Base introuvable. Lancez agent_documents.py d'abord.")

        st.divider()

        # Bouton nouvelle conversation
        if st.button("🗑️ Nouvelle conversation", use_container_width=True):
            for key in ["messages", "internal_history", "vision_history", "current_image_b64", "current_image_name"]:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()

        st.divider()

        # Afficher les documents créés
        st.subheader("📁 Documents créés")
        if os.path.exists(OUTPUTS_DIR):
            files = sorted(os.listdir(OUTPUTS_DIR))
            if files:
                for f in files:
                    fpath = os.path.join(OUTPUTS_DIR, f)
                    fsize = os.path.getsize(fpath)
                    st.caption(f"📄 {f} ({fsize:,} octets)")
            else:
                st.caption("Aucun document pour l'instant.")
        else:
            st.caption("Dossier outputs/ vide.")

    # ==================== INIT ÉTAT ====================
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "internal_history" not in st.session_state:
        st.session_state.internal_history = []
    if "vision_history" not in st.session_state:
        st.session_state.vision_history = []
    if "current_image_b64" not in st.session_state:
        st.session_state.current_image_b64 = None
    if "current_image_name" not in st.session_state:
        st.session_state.current_image_name = None
    if "doc_created" not in st.session_state:
        st.session_state.doc_created = False

    # Charger mémoire si demandé
    if use_memory and not st.session_state.messages:
        loaded = load_memory()
        if loaded:
            st.session_state.internal_history = loaded
            for i in range(0, len(loaded), 2):
                if i < len(loaded) and loaded[i].get("role") == "user":
                    st.session_state.messages.append({"role": "user", "content": loaded[i].get("content", "")})
                if i + 1 < len(loaded) and loaded[i+1].get("role") == "assistant":
                    st.session_state.messages.append({"role": "assistant", "content": loaded[i+1].get("content", "")})
            st.sidebar.info(f"Mémoire chargée : {len(loaded)} messages")

    # ==================== AFFICHAGE HISTORIQUE ====================
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            if msg.get("image_b64"):
                st.image(BytesIO(base64.b64decode(msg["image_b64"])), use_container_width=True)
            if msg.get("content"):
                st.markdown(msg["content"])

    # ==================== INPUT ====================
    user_input = st.chat_input("Posez votre question, décrivez une image, ou demandez un document...")

    # ==================== TRAITEMENT ====================
    if user_input or uploaded_file:
        if uploaded_file is not None:
            b64 = encode_image(uploaded_file)
            st.session_state.current_image_b64 = b64
            st.session_state.current_image_name = uploaded_file.name

        has_image = st.session_state.current_image_b64 is not None
        is_vision = has_image and (user_input or uploaded_file is not None)

        user_display_text = user_input if user_input else "Analyse cette image."
        with st.chat_message("user"):
            if has_image:
                st.image(BytesIO(base64.b64decode(st.session_state.current_image_b64)), use_container_width=True)
            st.markdown(user_display_text)

        st.session_state.messages.append({
            "role": "user",
            "content": user_display_text,
            "image_b64": st.session_state.current_image_b64 if has_image else None
        })

        with st.chat_message("assistant"):
            with st.spinner("Analyse en cours..."):
                try:
                    if is_vision:
                        reply = run_vision_agent(
                            st.session_state.vision_history,
                            user_display_text,
                            st.session_state.current_image_b64
                        )
                    else:
                        if len(st.session_state.internal_history) > MAX_TEXT_HISTORY:
                            st.session_state.internal_history = summarize_text_history(
                                st.session_state.internal_history
                            )
                            st.sidebar.info("Historique texte résumé.")
                        reply = run_text_agent(
                            st.session_state.internal_history,
                            user_display_text
                        )
                except Exception as e:
                    reply = f"❌ Erreur : {e}"

            st.markdown(reply)

            # Badge si un document a été créé
            if "Document créé avec succès" in reply or "créé avec succès" in reply:
                st.success("📄 Document généré ! Vérifiez la sidebar.")

        st.session_state.messages.append({"role": "assistant", "content": reply})

        if use_memory:
            save_memory(st.session_state.internal_history)


if __name__ == "__main__":
    main()
